import streamlit as st
import pytesseract
from pdf2image import convert_from_path
import tempfile
import PyPDF2
from io import BytesIO
from docx import Document
import time
from collections import Counter
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from textblob import TextBlob
from nltk.corpus import stopwords
import nltk

# Download stopwords for NLTK
nltk.download("stopwords")
stop_words = set(stopwords.words("english"))

# Set Streamlit theme using config
st.set_page_config(
    page_title="Intelligent Document Processing (IDP)",
    page_icon="🔧",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Title Section with Styling
st.markdown(
    """
    <h1 style="text-align: center; color: #4A96C9; font-size: 48px;">INTELLIGENT DOCUMENT PROCESSING (IDP)</h1>
    <p style="text-align: center; color: gray; font-size: 18px;">🔧 Upload PDFs to extract text, analyze content, and export results to Word!</p>
    """,
    unsafe_allow_html=True,
)

# Sidebar Customization
with st.sidebar:
    st.markdown("### Customize Appearance")
    theme_color = st.color_picker("Choose Theme Color", "#4CAF50")
    ocr_mode = st.radio("Text Extraction Mode", ["Auto Detect", "OCR Only", "Text-based Only"], index=0)
    st.markdown("---")
    st.markdown("### Additional Settings")
    keyword_analysis = st.checkbox("Perform Keyword Analysis", value=True)

st.write("\n")  # Add spacing

# File uploader
uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
if uploaded_file is not None:
    st.write(f"Uploaded file: **{uploaded_file.name}**")

    with st.spinner("Uploading file..."):
        time.sleep(1)  # Simulate upload animation

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        temp_file.write(uploaded_file.read())
        temp_pdf_path = temp_file.name

    try:
        extracted_text = ""

        with st.spinner("Extracting text from the PDF..."):
            time.sleep(1)  # Simulate processing animation
            pdf_reader = PyPDF2.PdfReader(temp_pdf_path)
            for page in pdf_reader.pages:
                extracted_text += page.extract_text()

        if not extracted_text.strip() or ocr_mode == "OCR Only":
            st.warning("No text found or OCR Mode selected. Using OCR...")

            with st.spinner("Performing OCR on PDF..."):
                images = convert_from_path(temp_pdf_path)
                ocr_text = ""
                for page_num, img in enumerate(images, start=1):
                    with st.spinner(f"Processing Page {page_num}..."):
                        time.sleep(0.5)  # Simulate page-by-page OCR animation
                        text = pytesseract.image_to_string(img)
                        ocr_text += f"\n--- Page {page_num} ---\n{text}"

                extracted_text = ocr_text

        if extracted_text.strip():
            # Display extracted text
            st.subheader("Extracted Text:")
            st.text_area("Text Output", extracted_text, height=300)

            # Word and Page Summary
            word_count = len(extracted_text.split())
            page_count = extracted_text.count("--- Page") or len(pdf_reader.pages)

            st.info(f"**Summary:** {word_count} words across {page_count} pages.")

            # Prepare data for Word export
            st.write("\nExport the extracted text to Word:")
            with st.spinner("Preparing Word document..."):
                time.sleep(1)  # Simulate document preparation animation
                doc = Document()

                if "--- Page" in extracted_text:  # For OCR-based text
                    for page_text in extracted_text.split("\n--- Page"):
                        if page_text.strip():
                            page_number, text = page_text.split("---\n", 1)
                            doc.add_heading(f"Page {page_number.strip()}", level=2)
                            doc.add_paragraph(text.strip())
                else:  # For text-based extraction
                    doc.add_paragraph(extracted_text.strip())

                output = BytesIO()
                doc.save(output)
                output.seek(0)

            st.download_button(
                label="📥 Download Extracted Text as Word",
                data=output,
                file_name="extracted_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            with st.success("Extraction and export completed successfully!"):
                time.sleep(1)  # Success animation

            # Keyword Analysis and Display
            if keyword_analysis:
                # Filter out stopwords and non-alphabetic words
                words = [
                    word.lower()
                    for word in extracted_text.split()
                    if word.isalpha() and word.lower() not in stop_words
                ]

                # Keyword Frequency Analysis
                keyword_counts = Counter(words)
                common_keywords = keyword_counts.most_common(10)

                # Display Top Keywords
                st.markdown("### Top 10 Keywords")
                for keyword, count in common_keywords:
                    st.write(f"**{keyword}:** {count} occurrences")

            # Word Cloud Visualization (optional)
            if extracted_text.strip():
                wordcloud = WordCloud(
                    width=800, height=400, background_color="white", stopwords=stop_words
                ).generate(extracted_text)

                plt.figure(figsize=(10, 5))
                plt.imshow(wordcloud, interpolation="bilinear")
                plt.axis("off")
                st.pyplot(plt)

            # Sentiment Analysis
            if extracted_text.strip():
                blob = TextBlob(extracted_text)
                sentiment = blob.sentiment
                st.write(
                    f"**Sentiment Analysis:** Polarity: {sentiment.polarity:.2f}, Subjectivity: {sentiment.subjectivity:.2f}"
                )

    except Exception as e:
        st.error(f"An error occurred: {e}")

# Footer Section
st.markdown(
    f"""
    <hr>
    <p style="text-align: center; color: {theme_color};">Developed with ❤️ using DataSparkLabs.</p>
    """,
    unsafe_allow_html=True,
)
